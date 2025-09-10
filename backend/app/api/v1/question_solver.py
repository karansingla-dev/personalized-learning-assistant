# backend/app/api/v1/question_solver.py
"""
AI Question Solver API with Direct Gemini Integration
Handles file uploads, question extraction, AI solving, and PDF generation
"""

from fastapi import APIRouter, UploadFile, File, HTTPException, Request, Form
from fastapi.responses import FileResponse
from typing import Optional, List, Dict
import traceback
import os
import io
import re
from datetime import datetime
import tempfile
from pathlib import Path
import json

# Document processing
import PyPDF2
from docx import Document as DocxDocument
from PIL import Image
import pytesseract  # For OCR
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
from reportlab.pdfgen import canvas
from dotenv import load_dotenv

# AI Integration - Direct Gemini import
import google.generativeai as genai
from app.config import settings
import logging

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/api/v1/question-solver", tags=["question-solver"])

# Constants
ALLOWED_EXTENSIONS = {'.pdf', '.docx', '.doc', '.png', '.jpg', '.jpeg'}
MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB

load_dotenv()

# ===== GEMINI CONFIGURATION =====
# Initialize Gemini directly in this file
gemini_model = None

def initialize_gemini():
    """Initialize Gemini model with fallback options"""
    global gemini_model
    
    api_key = settings.GEMINI_API_KEY or os.getenv('GEMINI_API_KEY')
    
    if not api_key:
        logger.warning("⚠️ GEMINI_API_KEY not found in environment variables")
        return None
    
    try:
        genai.configure(api_key=api_key)
        
        # Try different models in order of preference
        model_options = [
            'gemini-pro'
        ]
        
        for model_name in model_options:
            try:
                test_model = genai.GenerativeModel(model_name)
                # Test the model
                test_response = test_model.generate_content("Test")
                if test_response:
                    gemini_model = test_model
                    logger.info(f"✅ Successfully initialized Gemini model: {model_name}")
                    return gemini_model
            except Exception as e:
                logger.debug(f"Failed to initialize {model_name}: {e}")
                continue
        
        logger.error("❌ Could not initialize any Gemini model")
        return None
        
    except Exception as e:
        logger.error(f"❌ Failed to configure Gemini API: {e}")
        return None

# Initialize Gemini when module loads
# initialize_gemini()
model = genai.GenerativeModel('gemini-pro')

# ===== END GEMINI CONFIGURATION =====

# Supported file types
ALLOWED_EXTENSIONS = {'.pdf', '.docx', '.doc', '.png', '.jpg', '.jpeg'}
MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB

class QuestionExtractor:
    """Extract questions from various file formats"""
    
    @staticmethod
    async def extract_from_pdf(file_content: bytes) -> str:
        """Extract text from PDF"""
        try:
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            text = ""
            
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text += page.extract_text() + "\n"
            
            return text.strip()
        except Exception as e:
            logger.error(f"PDF extraction error: {e}")
            raise HTTPException(status_code=400, detail="Failed to extract text from PDF")
    
    @staticmethod
    async def extract_from_docx(file_content: bytes) -> str:
        """Extract text from DOCX"""
        try:
            doc = DocxDocument(io.BytesIO(file_content))
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text.strip()
        except Exception as e:
            logger.error(f"DOCX extraction error: {e}")
            raise HTTPException(status_code=400, detail="Failed to extract text from DOCX")
    
    @staticmethod
    async def extract_from_image(file_content: bytes) -> str:
        """Extract text from image using OCR"""
        try:
            # Open image
            image = Image.open(io.BytesIO(file_content))
            
            # Convert to RGB if necessary
            if image.mode != 'RGB':
                image = image.convert('RGB')
            
            # Perform OCR
            text = pytesseract.image_to_string(image, lang='eng')
            
            # Clean up the text
            text = re.sub(r'\n+', '\n', text)
            text = text.strip()
            
            if not text:
                raise ValueError("No text found in image")
            
            return text
        except Exception as e:
            logger.error(f"OCR extraction error: {e}")
            raise HTTPException(status_code=400, detail="Failed to extract text from image. Make sure the image is clear and contains readable text.")

class QuestionParser:
    """Enhanced parser for academic question papers (CBSE/ICSE format)"""
    
    @staticmethod
    def parse_questions(text: str) -> List[Dict]:
        """Identify and parse questions from academic papers"""
        
        questions = []
        
        # Clean the text first
        text = QuestionParser._clean_text(text)
        
        # Try multiple parsing strategies
        # Strategy 1: Parse numbered questions (most common in academic papers)
        numbered_questions = QuestionParser._parse_numbered_questions(text)
        if numbered_questions:
            questions.extend(numbered_questions)
        
        # Strategy 2: Parse sub-questions (like (A), (B), (i), (ii))
        sub_questions = QuestionParser._parse_sub_questions(text)
        if sub_questions:
            questions.extend(sub_questions)
        
        # Strategy 3: Parse case study questions
        case_questions = QuestionParser._parse_case_study_questions(text)
        if case_questions:
            questions.extend(case_questions)
        
        # Remove duplicates and clean
        unique_questions = QuestionParser._remove_duplicates(questions)
        
        # Limit to reasonable number
        return unique_questions[:50]  # Increased limit for full papers
    
    @staticmethod
    def _clean_text(text: str) -> str:
        """Clean and normalize text"""
        # Remove page numbers, headers, footers
        text = re.sub(r'Page \d+|page \d+|\d+\s*$', '', text, flags=re.MULTILINE)
        
        # Remove excessive whitespace
        text = re.sub(r'\n{3,}', '\n\n', text)
        text = re.sub(r' {2,}', ' ', text)
        
        # Fix common OCR issues
        text = text.replace('𝑜', '°')  # degree symbol
        text = text.replace('𝑂', '°')
        
        return text.strip()
    
    @staticmethod
    def _parse_numbered_questions(text: str) -> List[Dict]:
        """Parse standard numbered questions (1., 2., etc.)"""
        questions = []
        
        # Pattern for numbered questions with various formats
        patterns = [
            # Standard numbering: "1. Question text"
            r'(?:^|\n)(\d{1,2})\.\s+(.+?)(?=\n\d{1,2}\.|\nSection|\nOR\n|\nDIRECTION|\Z)',
            
            # Question with number in box or after Q: "Q1" or "Question 1:"
            r'(?:^|\n)(?:Q|Question)\s*(\d{1,2})[\.:]\s*(.+?)(?=\n(?:Q|Question)\s*\d{1,2}|\nSection|\Z)',
        ]
        
        for pattern in patterns:
            matches = re.findall(pattern, text, re.MULTILINE | re.DOTALL)
            
            for match in matches:
                if len(match) >= 2:
                    q_num = match[0]
                    q_text = match[1].strip()
                    
                    # Clean the question text
                    q_text = QuestionParser._clean_question_text(q_text)
                    
                    # Check if it's a valid question
                    if QuestionParser._is_valid_question(q_text):
                        # Check for MCQ options
                        mcq_data = QuestionParser._extract_mcq_options(q_text)
                        
                        questions.append({
                            'number': q_num,
                            'question': mcq_data['question'] if mcq_data else q_text,
                            'type': 'MCQ' if mcq_data else 'descriptive',
                            'options': mcq_data['options'] if mcq_data else [],
                            'marks': QuestionParser._extract_marks(q_text)
                        })
        
        return questions
    
    @staticmethod
    def _parse_sub_questions(text: str) -> List[Dict]:
        """Parse sub-questions like (A), (B), (i), (ii)"""
        questions = []
        
        # Pattern for sub-questions
        patterns = [
            r'\(([A-Za-z])\)\s*(.+?)(?=\([A-Za-z]\)|\nOR\n|\n\d+\.|\Z)',
            r'\(([ivxIVX]+)\)\s*(.+?)(?=\([ivxIVX]+\)|\n\d+\.|\Z)',
        ]
        
        for pattern in patterns:
            matches = re.findall(pattern, text, re.MULTILINE | re.DOTALL)
            
            for match in matches:
                if len(match) >= 2:
                    sub_label = match[0]
                    q_text = match[1].strip()
                    
                    q_text = QuestionParser._clean_question_text(q_text)
                    
                    if QuestionParser._is_valid_question(q_text):
                        questions.append({
                            'number': f'({sub_label})',
                            'question': q_text,
                            'type': 'sub-question',
                            'options': [],
                            'marks': QuestionParser._extract_marks(q_text)
                        })
        
        return questions
    
    @staticmethod
    def _parse_case_study_questions(text: str) -> List[Dict]:
        """Parse case study based questions"""
        questions = []
        
        # Look for case study indicators
        case_study_pattern = r'(?:case study|Case Study|CASE STUDY).+?(?=case study|Case Study|CASE STUDY|\Z)'
        case_studies = re.findall(case_study_pattern, text, re.IGNORECASE | re.DOTALL)
        
        for case in case_studies:
            # Extract questions from case study
            # These usually have (i), (ii), (iii) format
            sub_q_pattern = r'\(([ivx]+)\)\s*(.+?)(?=\([ivx]+\)|\Z)'
            matches = re.findall(sub_q_pattern, case, re.MULTILINE | re.DOTALL)
            
            for match in matches:
                if len(match) >= 2:
                    q_text = match[1].strip()
                    q_text = QuestionParser._clean_question_text(q_text)
                    
                    if QuestionParser._is_valid_question(q_text):
                        questions.append({
                            'number': f'Case-{match[0]}',
                            'question': q_text,
                            'type': 'case-study',
                            'options': [],
                            'marks': QuestionParser._extract_marks(q_text)
                        })
        
        return questions
    
    @staticmethod
    def _extract_mcq_options(text: str) -> Optional[Dict]:
        """Extract MCQ options from question text"""
        
        # Pattern for options A), B), C), D) or A. B. C. D.
        option_pattern = r'[A-D][)\.]?\s*(.+?)(?=[A-D][)\.]|\n\d+\.|\Z)'
        
        # Check if text contains MCQ options
        if re.search(r'[A-D][)\.]', text):
            # Split question and options
            parts = re.split(r'(?=[A-D][)\.])', text, maxsplit=1)
            
            if len(parts) >= 2:
                question_text = parts[0].strip()
                options_text = parts[1]
                
                # Extract individual options
                options = []
                option_matches = re.findall(option_pattern, options_text, re.MULTILINE)
                
                for opt in option_matches:
                    opt_text = opt.strip()
                    if opt_text and len(opt_text) > 0:
                        options.append(opt_text)
                
                if len(options) >= 2:  # Valid MCQ should have at least 2 options
                    return {
                        'question': question_text,
                        'options': options
                    }
        
        return None
    
    @staticmethod
    def _clean_question_text(text: str) -> str:
        """Clean individual question text"""
        
        # Remove marks notation
        text = re.sub(r'\[\d+\s*marks?\]|\(\d+\s*marks?\)|\d+\s*marks?$', '', text, flags=re.IGNORECASE)
        
        # Remove "OR" sections for now (can be handled separately)
        if '\nOR\n' in text:
            text = text.split('\nOR\n')[0]
        
        # Clean up whitespace
        text = ' '.join(text.split())
        
        return text.strip()
    
    @staticmethod
    def _is_valid_question(text: str) -> bool:
        """Check if text is a valid question"""
        
        # Minimum length check
        if len(text) < 10:
            return False
        
        # Skip section headers and instructions
        skip_patterns = [
            r'^Section [A-E]',
            r'^SECTION [A-E]',
            r'^General Instructions',
            r'^Instructions',
            r'^Note:',
            r'^Answer',
            r'consists of \d+ questions',
            r'marks each'
        ]
        
        for pattern in skip_patterns:
            if re.match(pattern, text, re.IGNORECASE):
                return False
        
        # Check for question indicators
        question_indicators = [
            '?',  # Has question mark
            r'\bfind\b', r'\bcalculate\b', r'\bprove\b', r'\bsolve\b',
            r'\bexplain\b', r'\bdefine\b', r'\bderive\b', r'\bshow that\b',
            r'\bwhat\b', r'\bwhich\b', r'\bhow\b', r'\bwhy\b', r'\bwhen\b',
            r'\bevaluate\b', r'\bdetermine\b', r'\bstate\b', r'\bwrite\b',
            r'\bif\b.*\bthen\b', r'\bgiven\b.*\bfind\b'
        ]
        
        text_lower = text.lower()
        for indicator in question_indicators:
            if '?' == indicator:
                if indicator in text:
                    return True
            elif re.search(indicator, text_lower):
                return True
        
        # Check if it looks like a math problem
        if re.search(r'[\d\+\-\*\/\=\^]+', text) and len(text) > 15:
            return True
        
        return False
    
    @staticmethod
    def _extract_marks(text: str) -> Optional[int]:
        """Extract marks from question text"""
        
        marks_pattern = r'\[(\d+)\s*marks?\]|\((\d+)\s*marks?\)|(\d+)\s*marks?'
        match = re.search(marks_pattern, text, re.IGNORECASE)
        
        if match:
            for group in match.groups():
                if group:
                    return int(group)
        
        return None
    
    @staticmethod
    def _remove_duplicates(questions: List[Dict]) -> List[Dict]:
        """Remove duplicate questions"""
        
        seen = set()
        unique = []
        
        for q in questions:
            # Create a unique key for the question
            q_text = q['question'][:100]  # First 100 chars
            q_key = re.sub(r'[^a-zA-Z0-9]', '', q_text.lower())
            
            if q_key not in seen and len(q_key) > 10:
                seen.add(q_key)
                unique.append(q)
        
        return unique


# Also add this helper function if it doesn't exist
def format_parsed_questions(questions: List[Dict]) -> List[Dict]:
    """Format parsed questions for solving"""
    
    formatted = []
    
    for i, q in enumerate(questions, 1):
        formatted_q = {
            'question': q.get('question', ''),
            'original_number': q.get('number', str(i)),
            'type': q.get('type', 'unknown'),
            'marks': q.get('marks'),
        }
        
        # For MCQs, include options in the question
        if q.get('options'):
            options_text = '\n'.join([f'{chr(65+j)}) {opt}' for j, opt in enumerate(q['options'])])
            formatted_q['question'] = f"{q['question']}\n\nOptions:\n{options_text}"
        
        formatted.append(formatted_q)
    
    return formatted

class QuestionSolver:
    """Enhanced solver for academic questions"""
    
    @staticmethod
    async def solve_questions_enhanced(
        questions: List[Dict], 
        subject: str = "Mathematics", 
        class_level: int = 10
    ) -> List[Dict]:
        """Generate AI solutions for academic questions"""
        
        solved_questions = []
        
        for i, q in enumerate(questions, 1):
            try:
                # Determine question type and adjust solving approach
                question_type = q.get('type', 'unknown')
                # FIX: Handle None marks properly
                marks = q.get('marks')
                if marks is None:
                    marks = 1  # Default to 1 mark
                
                solution = await QuestionSolver.solve_academic_question(
                    question=q['question'],
                    subject=subject,
                    class_level=class_level,
                    question_number=q.get('original_number', str(i)),
                    question_type=question_type,
                    marks=marks
                )
                
                solved_questions.append({
                    'number': q.get('original_number', str(i)),
                    'question': q['question'],
                    'type': question_type,
                    'marks': marks,
                    'solution': solution['solution'],
                    'steps': solution.get('steps', []),
                    'final_answer': solution.get('final_answer', ''),
                    'explanation': solution.get('explanation', '')
                })
                
            except Exception as e:
                logger.error(f"Error solving question {i}: {e}")
                solved_questions.append({
                    'number': q.get('original_number', str(i)),
                    'question': q['question'],
                    'type': q.get('type', 'unknown'),
                    'marks': q.get('marks', 1),
                    'solution': "Unable to generate solution for this question.",
                    'steps': [],
                    'final_answer': '',
                    'explanation': ''
                })
        
        return solved_questions

    @staticmethod
    async def solve_academic_question(
        question: str,
        subject: str,
        class_level: int,
        question_number: str,
        question_type: str = "unknown",
        marks: Optional[int] = None  # Make marks optional
    ) -> Dict:
        """Solve a single academic question with appropriate detail"""
        
        if not gemini_model:
            model = genai.GenerativeModel('gemini-pro')
            # initialize_gemini()
            
            if not gemini_model:
                return {
                    'solution': "AI service is not available.",
                    'steps': [],
                    'final_answer': '',
                    'explanation': ''
                }
        
        # FIX: Handle None marks properly
        if marks is None:
            marks = 1  # Default to 1 mark if not specified
        
        # Adjust prompt based on question type and marks
        detail_level = "detailed" if marks >= 3 else "concise"
        
        prompt = f"""
        You are an expert {subject} teacher helping a Class {class_level} student solve a question from their exam paper.
        
        Question Number: {question_number}
        Question Type: {question_type}
        Marks: {marks}
        
        Question: {question}
        
        Please provide a {detail_level} solution appropriate for {marks} marks.
        
        Guidelines:
        - For MCQ: Explain why the correct option is right and others are wrong
        - For 1-2 marks: Give a concise solution with key steps
        - For 3-5 marks: Provide detailed step-by-step solution with explanations
        - For mathematical problems: Show all calculations clearly
        - Use appropriate formulas and theorems
        - Follow CBSE/Academic answer writing pattern
        
        Format your response as JSON:
        {{
            "solution": "Complete solution text",
            "steps": [
                "Step 1: Understanding the problem...",
                "Step 2: Applying the formula/concept...",
                "Step 3: Calculation/Working...",
                ...
            ],
            "final_answer": "The final answer with proper units",
            "explanation": "Key concept or learning point"
        }}
        
        Make the solution clear and exam-appropriate for Class {class_level}.
        """
        
        try:
            response = gemini_model.generate_content(prompt)
            
            if not response or not response.text:
                return {
                    'solution': "Failed to generate solution.",
                    'steps': [],
                    'final_answer': '',
                    'explanation': ''
                }
            
            response_text = response.text
            
            # Clean and parse response
            if "```json" in response_text:
                response_text = response_text.split("```json")[1].split("```")[0]
            elif "```" in response_text:
                response_text = response_text.split("```")[1].split("```")[0]
            
            try:
                solution_data = json.loads(response_text)
                
                return {
                    'solution': solution_data.get('solution', 'Solution generated'),
                    'steps': solution_data.get('steps', []),
                    'final_answer': solution_data.get('final_answer', ''),
                    'explanation': solution_data.get('explanation', '')
                }
                
            except json.JSONDecodeError:
                # Fallback: format as structured text
                return {
                    'solution': response_text,
                    'steps': response_text.split('\n') if '\n' in response_text else [response_text],
                    'final_answer': '',
                    'explanation': ''
                }
                
        except Exception as e:
            logger.error(f"Error in AI solution generation: {e}")
            return {
                'solution': f"Error generating solution: {str(e)}",
                'steps': [],
                'final_answer': '',
                'explanation': ''
            }

# Enhanced PDF Generator for academic papers
class PDFGenerator:
    """Generate professional PDF for academic solutions"""
    
    @staticmethod
    def generate_enhanced_pdf(
        solved_questions: List[Dict], 
        student_name: str = "Student",
        subject: str = "Mathematics",
        class_level: int = 10
    ) -> bytes:
        """Create a professional PDF with academic formatting"""
        
        # Create a temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_file:
            pdf_path = tmp_file.name
        
        # Create PDF document
        doc = SimpleDocTemplate(
            pdf_path,
            pagesize=A4,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=18,
        )
        
        # Container for the 'Flowable' objects
        elements = []
        
        # Define styles
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Title'],
            fontSize=24,
            textColor=colors.HexColor('#1e3a8a'),
            spaceAfter=30,
            alignment=TA_CENTER
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            textColor=colors.HexColor('#7c3aed'),
            spaceAfter=12,
            spaceBefore=12,
            leftIndent=0
        )
        
        question_style = ParagraphStyle(
            'Question',
            parent=styles['Normal'],
            fontSize=12,
            textColor=colors.HexColor('#1f2937'),
            spaceAfter=6,
            leftIndent=20,
            bulletIndent=0
        )
        
        solution_style = ParagraphStyle(
            'Solution',
            parent=styles['Normal'],
            fontSize=11,
            textColor=colors.HexColor('#374151'),
            spaceAfter=6,
            leftIndent=20,
            rightIndent=20,
            alignment=TA_JUSTIFY
        )
        
        answer_style = ParagraphStyle(
            'Answer',
            parent=styles['Normal'],
            fontSize=12,
            textColor=colors.HexColor('#059669'),
            spaceAfter=12,
            leftIndent=20,
            fontName='Helvetica-Bold'
        )
        
        # Add title
        elements.append(Paragraph(f"{subject} Solutions - Class {class_level}", title_style))
        
        # Calculate total marks safely
        total_marks = sum(q.get('marks', 0) for q in solved_questions if q.get('marks') is not None)
        
        # Add metadata
        metadata = f"""
        <para align="center">
        <b>Student:</b> {student_name}<br/>
        <b>Date:</b> {datetime.now().strftime('%B %d, %Y')}<br/>
        <b>Total Questions:</b> {len(solved_questions)}<br/>
        <b>Total Marks:</b> {total_marks}
        </para>
        """
        elements.append(Paragraph(metadata, styles['Normal']))
        elements.append(Spacer(1, 0.5*inch))
        
        # Add questions and solutions
        for q in solved_questions:
            # Question number and text
            q_marks = q.get('marks', '')
            marks_text = f" [{q_marks} marks]" if q_marks else ""
            question_text = f"<b>Question {q['number']}{marks_text}:</b> {q['question']}"
            
            # Clean the text for XML
            question_text = question_text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            
            elements.append(Paragraph(question_text, question_style))
            elements.append(Spacer(1, 0.1*inch))
            
            # Solution heading
            elements.append(Paragraph("<b>Solution:</b>", heading_style))
            
            # Solution steps
            if q.get('steps') and len(q['steps']) > 0:
                for step in q['steps']:
                    # Clean step text for XML
                    step_text = f"• {step}".replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    elements.append(Paragraph(step_text, solution_style))
            else:
                # Clean solution text for XML
                sol_text = q.get('solution', 'Solution not available')
                sol_text = sol_text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                elements.append(Paragraph(sol_text, solution_style))
            
            # Final answer
            if q.get('final_answer'):
                # Clean answer text for XML
                ans_text = f"<b>Final Answer:</b> {q['final_answer']}"
                ans_text = ans_text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                elements.append(Paragraph(ans_text, answer_style))
            
            # Explanation if available
            if q.get('explanation'):
                elements.append(Paragraph("<b>Explanation:</b>", heading_style))
                # Clean explanation text for XML
                exp_text = q['explanation'].replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                elements.append(Paragraph(exp_text, solution_style))
            
            # Add separator
            elements.append(Spacer(1, 0.3*inch))
            
            # Line separator
            line_data = [[''] * 10]
            line_table = Table(line_data, colWidths=[0.6*inch]*10)
            line_table.setStyle(TableStyle([
                ('LINEBELOW', (0, 0), (-1, 0), 1, colors.HexColor('#e5e7eb')),
            ]))
            elements.append(line_table)
            elements.append(Spacer(1, 0.3*inch))
        
        # Build PDF
        doc.build(elements)
        
        # Read the PDF file
        with open(pdf_path, 'rb') as f:
            pdf_content = f.read()  # NOW pdf_content is defined!
        
        # Clean up temporary file
        os.unlink(pdf_path)
        
        return pdf_content  # NOW this will work!

@router.post("/upload-and-solve")
async def upload_and_solve(
    request: Request,
    file: UploadFile = File(...),
    subject: str = Form("Mathematics"),
    class_level: int = Form(10),
    student_name: str = Form("Student")
):
    """
    Upload a file with questions and get AI solutions
    Enhanced for CBSE/Academic paper format
    """
    
    # DEFINE file_ext HERE - This was missing!
    file_ext = Path(file.filename).suffix.lower()
    
    # Validate file extension
    if file_ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"File type {file_ext} not supported. Allowed types: {', '.join(ALLOWED_EXTENSIONS)}"
        )
    
    # Read file content
    file_content = await file.read()
    
    # Check file size
    if len(file_content) > MAX_FILE_SIZE:
        raise HTTPException(status_code=400, detail="File size exceeds 10MB limit")
    
    try:
        # Extract text based on file type
        extractor = QuestionExtractor()
        
        if file_ext == '.pdf':
            extracted_text = await extractor.extract_from_pdf(file_content)
        elif file_ext in ['.docx', '.doc']:
            extracted_text = await extractor.extract_from_docx(file_content)
        elif file_ext in ['.png', '.jpg', '.jpeg']:
            extracted_text = await extractor.extract_from_image(file_content)
        else:
            raise HTTPException(status_code=400, detail="Unsupported file type")
        
        if not extracted_text:
            raise HTTPException(status_code=400, detail="No text found in the uploaded file")
        
        # Log extracted text length for debugging
        logger.info(f"Extracted text length: {len(extracted_text)} characters")
        
        # Parse questions with improved parser
        parser = QuestionParser()
        parsed_questions = parser.parse_questions(extracted_text)
        
        if not parsed_questions:
            # If no questions found, provide helpful error
            logger.warning("No questions found with standard parsing")
            raise HTTPException(
                status_code=400, 
                detail="No questions found. Please ensure the document contains numbered questions (1., 2., etc.) or clear question text."
            )
        
        logger.info(f"Successfully parsed {len(parsed_questions)} questions")
        
        # Format questions for solving (if you have this function)
        # If not, use parsed_questions directly
        try:
            questions_to_solve = format_parsed_questions(parsed_questions)
        except NameError:
            # If format_parsed_questions doesn't exist, use parsed questions directly
            questions_to_solve = []
            for i, q in enumerate(parsed_questions, 1):
                formatted_q = {
                    'question': q.get('question', ''),
                    'original_number': q.get('number', str(i)),
                    'type': q.get('type', 'unknown'),
                    'marks': q.get('marks'),
                }
                # For MCQs, include options in the question
                if q.get('options'):
                    options_text = '\n'.join([f'{chr(65+j)}) {opt}' for j, opt in enumerate(q['options'])])
                    formatted_q['question'] = f"{q['question']}\n\nOptions:\n{options_text}"
                questions_to_solve.append(formatted_q)
        
        # Log sample of parsed questions for debugging
        if questions_to_solve:
            logger.info(f"Sample question: {questions_to_solve[0]['question'][:100]}...")
        
        # Solve questions using AI
        solver = QuestionSolver()
        
        # Try enhanced solver first, fallback to regular if not available
        try:
            solved_questions = await solver.solve_questions_enhanced(
                questions_to_solve, 
                subject, 
                class_level
            )
        except AttributeError:
            # Fallback to regular solve_questions if enhanced version doesn't exist
            solved_questions = await solver.solve_questions(
                questions_to_solve, 
                subject, 
                class_level
            )
        
        # Generate PDF
        pdf_generator = PDFGenerator()
        
        # Try enhanced PDF generator, fallback to regular if not available
        try:
            pdf_content = pdf_generator.generate_enhanced_pdf(
                solved_questions, 
                student_name, 
                subject,
                class_level
            )
        except AttributeError:
            # Fallback to regular generate_pdf
            pdf_content = pdf_generator.generate_pdf(
                solved_questions, 
                student_name, 
                subject
            )
        
        # Save PDF temporarily
        temp_pdf_path = f"/tmp/solutions_{datetime.now().timestamp()}.pdf"
        with open(temp_pdf_path, 'wb') as f:
            f.write(pdf_content)
        
        # Return detailed response
        response_data = {
            "success": True,
            "message": f"Successfully solved {len(solved_questions)} questions from your {subject} paper",
            "questions_found": len(parsed_questions),
            "questions_solved": len(solved_questions),
            "pdf_path": temp_pdf_path,
            "solved_questions": solved_questions
        }
        
        # Add parsing details if available
        try:
            response_data["parsing_details"] = {
                "mcq_count": len([q for q in parsed_questions if q.get('type') == 'MCQ']),
                "descriptive_count": len([q for q in parsed_questions if q.get('type') == 'descriptive']),
                "total_marks": sum(q.get('marks', 0) for q in parsed_questions if q.get('marks'))
            }
        except:
            pass  # Parsing details are optional
        
        return response_data
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error processing file: {e}")
        logger.error(f"Traceback: {traceback.format_exc()}")  # Add detailed error logging
        raise HTTPException(status_code=500, detail=f"Error processing file: {str(e)}")

@router.get("/download-pdf/{filename}")
async def download_pdf(filename: str):
    """Download the generated PDF"""
    
    file_path = f"/tmp/{filename}"
    
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="File not found")
    
    return FileResponse(
        path=file_path,
        media_type='application/pdf',
        filename=f"solutions_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
    )

# Additional endpoint for solving single question
@router.post("/solve-single")
async def solve_single_question(
    request: Request,
    question: str = Form(...),
    subject: str = Form("Mathematics"),
    class_level: int = Form(10)
):
    """Solve a single question"""
    
    solver = QuestionSolver()
    solution = await solver.solve_single_question(question, subject, class_level)
    
    return {
        "success": True,
        "question": question,
        "solution": solution
    }